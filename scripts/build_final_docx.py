from pathlib import Path
import argparse
import re
import pyphen
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

INLINE_RE = re.compile(r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*))')
CH_RE = re.compile(r'^##\s+Kapitel\s+(\d+)\s+–\s+(.+?)\s*$')
GERMAN_WORD_RE = re.compile(r'[A-Za-zÄÖÜäöüß]{7,}')
HYPH = pyphen.Pyphen(lang='de_DE')


def set_doc_bool(doc, tag, enabled=True):
    settings = doc.settings._element
    node = settings.find(qn(f'w:{tag}'))
    if enabled:
        if node is None:
            node = OxmlElement(f'w:{tag}')
            settings.append(node)
        node.set(qn('w:val'), 'true')
    elif node is not None:
        settings.remove(node)


def set_doc_val(doc, tag, value):
    settings = doc.settings._element
    node = settings.find(qn(f'w:{tag}'))
    if node is None:
        node = OxmlElement(f'w:{tag}')
        settings.append(node)
    node.set(qn('w:val'), str(value))


def set_font_run(run, name='Garamond', size=12.5):
    run.font.name = name
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn(f'w:{attr}'), name)
    lang = rpr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rpr.append(lang)
    for attr in ('val', 'eastAsia', 'bidi'):
        lang.set(qn(f'w:{attr}'), 'de-DE')


def set_font_style(style, name, size, bold=None):
    style.font.name = name
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn(f'w:{attr}'), name)


def configure_section(section):
    section.page_width = Cm(12.85)
    section.page_height = Cm(19.84)
    section.top_margin = Cm(1.22)
    section.bottom_margin = Cm(1.22)
    section.left_margin = Cm(1.95)
    section.right_margin = Cm(1.40)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)


def add_page_number_field(paragraph, align):
    paragraph.clear()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), ' PAGE ')
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        fonts.set(qn(f'w:{attr}'), 'Garamond')
    size = OxmlElement('w:sz')
    size.set(qn('w:val'), '18')
    rpr.extend([fonts, size])
    run.append(rpr)
    text = OxmlElement('w:t')
    text.text = '1'
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def set_page_start(section, number=1):
    sect = section._sectPr
    node = sect.find(qn('w:pgNumType'))
    if node is None:
        node = OxmlElement('w:pgNumType')
        sect.append(node)
    node.set(qn('w:start'), str(number))


def add_toc_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-1" \\h \\z \\u '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    cached = OxmlElement('w:t')
    cached.text = 'Inhaltsverzeichnis aktualisieren.'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, cached, end])


def append_text(run_element, text):
    if not text:
        return
    node = OxmlElement('w:t')
    if text[0].isspace() or text[-1].isspace():
        node.set(qn('xml:space'), 'preserve')
    node.text = text
    run_element.append(node)


def soft_hyphenate_run(run):
    text = run.text.replace('\u00ad', '')
    element = run._r
    for child in list(element):
        if child.tag != qn('w:rPr'):
            element.remove(child)
    pos = 0
    for match in GERMAN_WORD_RE.finditer(text):
        append_text(element, text[pos:match.start()])
        word = match.group(0)
        cuts = [] if word.isupper() else HYPH.positions(word)
        last = 0
        for cut in cuts:
            append_text(element, word[last:cut])
            element.append(OxmlElement('w:softHyphen'))
            last = cut
        append_text(element, word[last:])
        pos = match.end()
    append_text(element, text[pos:])


def normalize_text(text):
    return text.replace('„', '»').replace('“', '«').replace('—', '–')


def add_inline(paragraph, text):
    text = normalize_text(text)
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_font_run(run)
        token = match.group(0)
        run = paragraph.add_run()
        if token.startswith('***') and token.endswith('***'):
            run.text = token[3:-3]
            run.bold = True
            run.italic = True
        elif token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
        else:
            run.text = token[1:-1]
            run.italic = True
        set_font_run(run)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font_run(run)
    for run in paragraph.runs:
        soft_hyphenate_run(run)


def parse(text):
    chapters = []
    current = None
    blocks = []
    paragraph_lines = []

    def flush_paragraph():
        nonlocal paragraph_lines
        if current is not None and paragraph_lines:
            value = ' '.join(line.strip() for line in paragraph_lines if line.strip()).strip()
            if value:
                blocks.append(('paragraph', value))
        paragraph_lines = []

    def flush_chapter():
        nonlocal current, blocks
        flush_paragraph()
        if current is not None:
            chapters.append((current, blocks))
        blocks = []

    for raw in text.splitlines():
        stripped = raw.strip()
        match = CH_RE.match(stripped)
        if match:
            flush_chapter()
            current = (int(match.group(1)), match.group(2))
            continue
        if current is None:
            continue
        if not stripped:
            flush_paragraph()
            continue
        if stripped == '---':
            flush_paragraph()
            if not blocks or blocks[-1][0] != 'scene_break':
                blocks.append(('scene_break', '*'))
            continue
        if stripped.startswith('#'):
            continue
        paragraph_lines.append(raw)

    flush_chapter()
    numbers = [meta[0] for meta, _ in chapters]
    if numbers != list(range(1, 41)):
        raise SystemExit(f'Chapter sequence mismatch: {numbers}')
    return chapters


def build(source: Path, output: Path):
    chapters = parse(source.read_text(encoding='utf-8'))
    doc = Document()
    doc.core_properties.title = 'ABWEICHUNG'
    doc.core_properties.subject = 'ABWEICHUNG – KDP Paperback 5.06 x 7.81 in'
    doc.core_properties.author = ''
    doc.core_properties.keywords = ''
    doc.core_properties.comments = ''

    set_doc_bool(doc, 'updateFields', True)
    set_doc_bool(doc, 'autoHyphenation', True)
    set_doc_bool(doc, 'doNotHyphenateCaps', True)
    set_doc_val(doc, 'hyphenationZone', 230)
    set_doc_val(doc, 'consecutiveHyphenLimit', 2)
    set_doc_bool(doc, 'mirrorMargins', True)
    set_doc_bool(doc, 'evenAndOddHeaders', True)

    first = doc.sections[0]
    configure_section(first)
    first.different_first_page_header_footer = True

    normal = doc.styles['Normal']
    set_font_style(normal, 'Garamond', 12.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0)
    normal.paragraph_format.widow_control = False

    heading = doc.styles['Heading 1']
    set_font_style(heading, 'Garamond', 14.5, True)
    heading.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = Cm(0)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(14)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True
    heading.paragraph_format.page_break_before = True

    names = [style.name for style in doc.styles]
    front = doc.styles.add_style('Front Matter', WD_STYLE_TYPE.PARAGRAPH) if 'Front Matter' not in names else doc.styles['Front Matter']
    set_font_style(front, 'Garamond', 12.5)
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.first_line_indent = Cm(0)

    scene = doc.styles.add_style('Scene Break', WD_STYLE_TYPE.PARAGRAPH) if 'Scene Break' not in names else doc.styles['Scene Break']
    set_font_style(scene, 'Garamond', 11.5)
    scene.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scene.paragraph_format.first_line_indent = Cm(0)
    scene.paragraph_format.space_before = Pt(8)
    scene.paragraph_format.space_after = Pt(8)
    scene.paragraph_format.keep_with_next = True

    for _ in range(5):
        doc.add_paragraph(style='Front Matter')
    p = doc.add_paragraph(style='Front Matter')
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('ABWEICHUNG')
    r.bold = True
    set_font_run(r, size=24)
    p = doc.add_paragraph(style='Front Matter')
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run('Wenn die Maschine recht hat')
    r.italic = True
    set_font_run(r, size=13)
    p = doc.add_paragraph(style='Front Matter')
    r = p.add_run('Du darfst widersprechen.')
    r.bold = True
    set_font_run(r, size=12)
    p = doc.add_paragraph(style='Front Matter')
    r = p.add_run('Die Beweislast liegt bei dir.')
    r.bold = True
    set_font_run(r, size=12)

    doc.add_page_break()
    doc.add_page_break()

    p = doc.add_paragraph(style='Front Matter')
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run('INHALT')
    r.bold = True
    set_font_run(r, size=18)
    toc = doc.add_paragraph(style='Front Matter')
    toc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_toc_field(toc)

    story = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(story)
    story.different_first_page_header_footer = False
    story.footer.is_linked_to_previous = False
    story.even_page_footer.is_linked_to_previous = False
    story.first_page_footer.is_linked_to_previous = False
    set_page_start(story, 1)
    add_page_number_field(story.footer.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_number_field(story.even_page_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.LEFT)

    for (number, title), blocks in chapters:
        heading_p = doc.add_paragraph(style='Heading 1')
        run = heading_p.add_run(f'Kapitel {number} – {title}')
        run.bold = True
        set_font_run(run, size=14.5)
        for kind, value in blocks:
            if kind == 'scene_break':
                p = doc.add_paragraph(style='Scene Break')
                run = p.add_run('*')
                set_font_run(run, size=11.5)
            else:
                p = doc.add_paragraph(style='Normal')
                add_inline(p, value)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    if output.stat().st_size < 180000:
        raise SystemExit(f'DOCX unexpectedly small: {output.stat().st_size}')
    print(f'Built {output} ({output.stat().st_size} bytes)')


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--source', default='ABWEICHUNG_FINAL.md')
    parser.add_argument('--output', default='ABWEICHUNG_FINAL.docx')
    args = parser.parse_args()
    build(Path(args.source), Path(args.output))


if __name__ == '__main__':
    main()
